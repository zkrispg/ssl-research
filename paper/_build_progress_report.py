"""Generate a Word progress report (.docx) for the SELD geometry-prior project.

Produces D:/ssl-research/paper/progress_report_2026-05-20.docx with:
  - Cover info (title, date, author)
  - 项目概况 / 研究问题
  - 已完成实验汇总（表）
  - 关键结果（表 + 段落）
  - 发现的问题与风险
  - 当前 Path C 进展（含 DCASE 2024 baseline 复现）
  - 下一步计划
  - 工作量盘点 / 附录
"""
from __future__ import annotations

import datetime as _dt
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_PATH = Path("D:/ssl-research/paper/progress_report_2026-05-20.docx")


def _set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 10.5,
                   align=None, color: RGBColor | None = None) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # type: ignore
    if color is not None:
        run.font.color.rgb = color


def _add_para(doc: Document, text: str, *, style: str | None = None,
              bold: bool = False, size_pt: float = 10.5,
              align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # type: ignore
    p.paragraph_format.space_after = Pt(4)


def _add_bullet(doc: Document, items: list[str], *, size_pt: float = 10.5) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(size_pt)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # type: ignore
        p.paragraph_format.space_after = Pt(2)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]],
               highlight_rows: set[int] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[j], h,
                       bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            color = None
            if highlight_rows and i in highlight_rows:
                color = RGBColor(0x9E, 0x2A, 0x2B)  # dark red emphasis
            _set_cell_text(
                table.rows[i + 1].cells[j], val,
                size_pt=10,
                align=WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT,
                color=color,
            )


def main() -> None:
    doc = Document()

    # ----- 文档默认字体 -----
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    rpr = style.element.rPr
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # type: ignore

    # ===================== 封面 =====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SELD 几何先验研究 · 实验进度汇报")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # type: ignore

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Geometry-Aware Attention for Sound Event Localization "
                      "and Detection: A Multi-Seed Empirical Study")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.name = "Times New Roman"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = _dt.date.today().isoformat()
    run = meta.add_run(f"汇报日期：{today}    研究方向：SELD / 空间音频 / 深度学习")
    run.font.size = Pt(10.5)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # type: ignore

    doc.add_paragraph()

    # ===================== 1. 项目概况 =====================
    doc.add_heading("1. 项目概况与研究问题", level=1)

    _add_para(doc,
              "本项目研究**几何感知注意力（Geometry-Aware Channel Attention, GCA）"
              "是否能为基于真实数据的 Sound Event Localization and Detection "
              "(SELD) 任务带来稳定的性能提升**。直觉上，把 4-mic 阵列的几何拓扑"
              "（成对间距、传声器位置）显式注入注意力模块应当帮助网络更高效地学习"
              "DOA 估计；但目前没有公开工作在多 seed、配对统计意义上系统验证过这点。"
              "我们以 STARSS23 真实数据为基础，做严格的多种子配对实验。")

    _add_para(doc, "核心研究问题（RQ）：", bold=True)
    _add_bullet(doc, [
        "RQ1：在控制训练预算的条件下，引入几何先验是否对 SELD 主指标（F1 / SELD score）"
        "带来配对显著的提升？",
        "RQ2：这一结论是否对模型容量（xs/m/l/xl）、SpecAugment 强度、"
        "音频格式（MIC vs FOA）、跨数据集（STARSS22）保持 robust？",
        "RQ3：在和 DCASE 2024 EUSIPCO 现代 baseline 同等训练 recipe 下，结论是否仍稳定？",
    ])

    _add_para(doc, "目标投稿会议：ICASSP 2027 / INTERSPEECH 2026 / TASLP（视实验深度而定）。", bold=False)

    # ===================== 2. 已完成实验汇总 =====================
    doc.add_heading("2. 已完成实验汇总", level=1)

    _add_para(doc,
              "截至本报告日，所有 Phase A（ICASSP-solid）+ Phase B（INTERSPEECH/TASLP "
              "扩展）的核心实验均已完成。Phase C（复现 DCASE 2024 现代 baseline）"
              "正在进行中。")

    headers = ["阶段", "实验内容", "训练 cell 数", "状态"]
    rows = [
        ["Phase A · 主对比",
         "no_geom vs full vs SELDnet（vanilla 30 ep）N=5 seeds",
         "15", "完成"],
        ["Phase A · 容量消融",
         "Capacity sweep at xs / m / l / xl × {no_geom, full} × N=5",
         "32 (含基础 m)", "完成"],
        ["Phase A · 增强对照",
         "SpecAugment {strong / weak / off} × {no_geom, full, SELDnet} × N=5",
         "20", "完成"],
        ["Phase B · 模态消融",
         "SELDnet · FOA × N=5（与 MIC 配对对比）",
         "5", "完成"],
        ["Phase B · 跨数据集",
         "STARSS23→STARSS22 dev-test zero-shot eval（15 ckpts）",
         "—", "完成"],
        ["Phase B · 统计强化",
         "Bootstrap CI、Cohen's d_z、power、Bonferroni、Wilcoxon",
         "—", "完成"],
        ["Phase C · 现代 baseline",
         "DCASE 2024 (EUSIPCO 2024) baseline 复现（FOA Multi-ACCDDOA）",
         "1（运行中）+ 4（待）", "进行中"],
    ]
    _add_table(doc, headers, rows)
    _add_para(doc, "合计：75 个主训练 cell + 1 个 Path C cell（运行至 epoch 32/60）。", bold=True)

    # ===================== 3. 关键结果 =====================
    doc.add_heading("3. 关键结果", level=1)

    doc.add_heading("3.1 主结论：几何先验在 N=5 paired t-test 下不显著（NS）", level=2)

    _add_para(doc, "在所有模型容量下，no_geom 与 full（带 GCA）的配对 t-test 均不显著：", bold=False)
    headers = ["容量", "参数量", "thr", "no_geom F1", "full F1", "Δ", "p"]
    rows = [
        ["xs",  "≈ 250 K",   "0.30", "6.55 %",  "6.55 %",  "+0.0",   "0.74"],
        ["m",   "≈ 590 K",   "0.30", "6.55 %",  "6.51 %",  "−0.04",  "0.84"],
        ["l",   "≈ 1.20 M",  "0.30", "—",       "—",       "−0.015", "0.69"],
        ["xl",  "≈ 2.31 M",  "0.30", "8.26 %",  "7.24 %",  "−0.99",  "0.79"],
    ]
    _add_table(doc, headers, rows)
    _add_para(doc, "解读：在 4 种容量、2 种阈值（0.18 / 0.30）下，所有 macro SELD 与 micro F1 "
                   "的配对 t-test p > 0.55；几何先验未能 detect 到稳定增益。", bold=False)

    doc.add_heading("3.2 SpecAugment 控制：弱/强 SpecAug 在 30 epoch 下都伤害性能", level=2)
    _add_bullet(doc, [
        "vanilla：6.5 % macro F1 / 14.9 % micro F1",
        "weak SpecAug：4.5 % / 12.0 %",
        "strong SpecAug：1.2 % / 3.4 %（崩塌）",
        "→ 30 epoch 训练预算下 SpecAug 还来不及补偿其造成的信息损失",
    ])

    doc.add_heading("3.3 SELDnet MIC vs FOA：MIC 在我们 pipeline 下显著更优", level=2)
    _add_para(doc, "（注意：与 EUSIPCO 2024 论文报的 FOA > MIC 相反，怀疑我们 IV 实现有 bug）",
              bold=False)
    headers = ["阈值", "指标", "MIC", "FOA", "Δ", "p"]
    rows = [
        ["0.30", "macro SELD", "0.969", "0.988", "+0.020", "0.016 *"],
        ["0.30", "micro F1",   "15.3 %", "4.4 %", "−10.9 %", "0.002 **"],
        ["0.18", "micro F1",   "13.5 %", "5.1 %", "−8.4 %",  "0.001 **"],
    ]
    _add_table(doc, headers, rows, highlight_rows={1, 2})

    doc.add_heading("3.4 跨数据集（STARSS23→STARSS22 zero-shot）：null 结论 robust", level=2)
    _add_bullet(doc, [
        "no_geom vs full @ thr=0.18：Δ macro SELD = +0.023，p = 0.62（NS）",
        "no_geom vs full @ thr=0.30：Δ macro SELD = −0.005，p = 0.64（NS）",
        "SELDnet macro SELD = 0.98（与 STARSS23 同水平，性能没崩）",
    ])

    doc.add_heading("3.5 Per-class Bonferroni（13 类 × 4 metric）", level=2)
    _add_bullet(doc, [
        "0 / 13 类通过 per-metric Bonferroni（α = 0.0038）",
        "0 / 13 类通过 global Bonferroni（α = 0.0010）",
        "Wilcoxon signed-rank 与 paired-t 结论一致",
        "→ 没有任何单一类别上 GCA 显示出稳定改进，主结论 null 全面 robust",
    ])

    # ===================== 4. 发现的问题 =====================
    doc.add_heading("4. 内审发现的问题与风险", level=1)
    _add_para(doc,
              "在确认上述 null 结论后，对照 ICASSP / INTERSPEECH 投稿标准做内审，"
              "识别出以下需要关注的风险：", bold=False)

    headers = ["风险", "等级", "说明", "应对"]
    rows = [
        ["绝对性能弱于 SOTA",
         "高",
         "我们的 SELDnet 复现在 30 epoch、vanilla recipe 下 macro F1 ≈ 6.5 %，"
         "DCASE 2024 baseline 报 F 20° = 13.1 %，top systems 达 54 %。"
         "Reviewer 会质疑 ablation 是否在弱 base 上做的。",
         "Path C：复现 EUSIPCO 2024 baseline，并在其上重做 GCA ablation。"],
        ["FOA 实现存疑",
         "中",
         "我们 FOA F1 < MIC F1，与官方 baseline 报的 FOA > MIC 相反，"
         "怀疑我们的 Intensity Vector 实现或归一化有 bug。",
         "Path C 包含官方 FOA 复现，可作 reference 排错。"],
        ["Cross-dataset 仅 STARSS22",
         "低",
         "对 INTERSPEECH 够用，对 TASLP 偏单薄。",
         "如必要可加 DCASE synthetic / L3DAS22 等。"],
        ["Distance estimation 未建模",
         "中",
         "DCASE 2024 metric 含距离 RDE，我们当前模型不输出距离。",
         "在 Path C 复现里包含距离 head（Multi-ACCDDOA 自带）。"],
    ]
    _add_table(doc, headers, rows)

    # ===================== 5. Path C 进展 =====================
    doc.add_heading("5. 当前进展：Path C — 复现 DCASE 2024 现代 baseline", level=1)

    doc.add_heading("5.1 论文研读", level=2)
    _add_para(doc,
              "已阅读并整理 Krause, Politis, Mesaros: "
              "“Sound Event Detection and Localization with Distance Estimation”，"
              "EUSIPCO 2024（DOI: 10.23919/EUSIPCO63174.2024.10715220）。"
              "完整研读笔记位于 paper/reading_notes_krause2024_eusipco.md。", bold=False)

    _add_para(doc, "该论文核心贡献：", bold=True)
    _add_bullet(doc, [
        "提出 Multi-ACCDDOA 表示，把 (x, y, z) DOA 向量扩成 (x, y, z, distance)",
        "ADPIT loss + 6 × 4 损失 / 范式对照（MT vs Multi-ACCDDOA × MSE/MAE/MSPE/MAPE）",
        "FOA 上 F₁ = 44.2 %（论文 1 秒段 + 2023 metric）"
        " / F 20° = 13.1 %（仓库帧级 + 2024 metric）",
        "MHSA × 2 块（来自 Sudarsanam et al. DCASE 2021）—— 这正是我们 GCA 想替换的对照",
    ])

    doc.add_heading("5.2 工程实施进展", level=2)
    _add_bullet(doc, [
        "Clone partha2409/DCASE2024_seld_baseline 到本地",
        "venv 适配：安装 IPython，stub 视频依赖（cv2/torchvision/PIL）",
        "修复 5 处 NumPy 2.0 兼容（np.NaN → np.nan）+ 3 处 np.int/np.float deprecation",
        "在 parameters.py 添加 task 100/101/102（FOA 60ep、smoke、MIC-GCC 60ep）",
        "给 train_seldnet.py 加 seed 参数（4th argv 或 SSL_SEED 环境变量）",
        "用 Windows junction 把 STARSS23 数据接入官方仓库要求的目录结构",
        "完成 4 分钟 FOA 特征抽取（168 wavs → 9.6 GB log-mel + IV + ADPIT labels）",
        "2-epoch smoke：F=3.2 %, AE=29°, RDE=0.28（pretrained synthetic init 生效）",
    ])

    doc.add_heading("5.3 实时复现进度（FOA seed=0, 60-epoch）", level=2)
    _add_para(doc, "截至本报告，正在跑 epoch 32/60，best epoch = 18：", bold=False)
    headers = ["指标", "我们 best epoch (epoch 18)", "EUSIPCO 2024 README 目标", "Δ vs 目标"]
    rows = [
        ["F 20°",         "12.0 %", "13.1 %", "−1.1 pp"],
        ["DOAE (CD)",     "32.7°",  "36.9°",  "−4.2°（更好）"],
        ["RDE",           "0.31",   "0.33",   "−0.02（更好）"],
        ["SELD score",    "0.57",   "≈ 0.58", "持平"],
    ]
    _add_table(doc, headers, rows)
    _add_para(doc,
              "结论：现在已在 baseline 性能附近。F 20° 略低 1.1 pp 但 DOAE 反而比官方更好，"
              "可能因为我们 batch=32（官方 128）+ patience 系统不同，或 jackknife 协议差异。"
              "等 60 epoch 跑完做最终对比。", bold=False)

    # ===================== 6. 下一步计划 =====================
    doc.add_heading("6. 下一步计划", level=1)
    headers = ["步骤", "任务", "预计 GPU 时长", "状态"]
    rows = [
        ["6.1", "完成 FOA seed=0 60-epoch（剩约 1 h）", "1 h", "进行中"],
        ["6.2", "扩 FOA × 4 seeds（seed=1..4）", "8 h", "待"],
        ["6.3", "MIC-GCC × 5 seeds", "10 h", "待"],
        ["6.4", "对照官方报的 13.1 % / 9.9 %，验证复现质量",
         "—", "待"],
        ["6.5", "把 GCA / no_geom / full 几何先验模块移植到 EUSIPCO baseline 架构",
         "（实现 1 d）", "待"],
        ["6.6", "在新 strong base 上跑 no_geom_modern × 5 + full_modern × 5", "10 h", "待"],
        ["6.7", "重做 paired t-test + bootstrap + per-class，对比新旧 base",
         "—", "待"],
        ["6.8", "更新论文初稿（Section 3 ↔ Section 4 ↔ Section 5）",
         "（文档 1 周）", "待"],
    ]
    _add_table(doc, headers, rows)
    _add_para(doc, "估算总 GPU 工时：约 30 h（不含未来 ablation 扩展）。", bold=True)

    # ===================== 7. 工作量盘点 =====================
    doc.add_heading("7. 工作量盘点（截至 2026-05-20）", level=1)
    _add_bullet(doc, [
        "75 个主训练 cell + 1 个进行中的 Path C cell（合计 76 cells）",
        "14 个 N=5 paired t-test JSON",
        "20 个 bootstrap + power JSON",
        "1 套 per-class Bonferroni + Wilcoxon 分析",
        "1 个 STARSS22 跨数据集 zero-shot eval（15 ckpts）",
        "1 份 Path C 仓库适配 + smoke 通过 + 60-epoch 启动",
        "1 份 EUSIPCO 2024 论文研读笔记",
        "1 份 2024 DCASE Task 3 SOTA 综述表（paper/survey_table_2024_seld.md）",
    ])

    # ===================== 8. 附录 =====================
    doc.add_heading("附录 A：工件位置", level=1)
    _add_bullet(doc, [
        "训练日志：D:\\ssl-research\\week11_starss23\\runs\\",
        "分析结果：runs\\analysis_*.json + runs\\multiseed_summary_*.json",
        "配对 t-test：runs\\pairwise_*_n5.json",
        "Path C 训练日志：runs\\dcase2024_repro_foa_seed0.log",
        "DCASE baseline 仓库：D:\\ssl-research\\dcase2024_baseline\\",
        "DCASE 数据 junction：D:\\ssl-research\\DCASE2024_SELD_dataset\\",
        "论文相关材料：D:\\ssl-research\\paper\\",
        "完整 transcript：D:\\ssl-research\\SESSION_SUMMARY.md",
    ])

    doc.add_heading("附录 B：硬件与环境", level=1)
    _add_bullet(doc, [
        "GPU：NVIDIA GeForce RTX 3050 Ti Laptop（4 GB VRAM）",
        "Python 3.10.6 / PyTorch 2.6.0+cu124 / CUDA 12.4",
        "OS：Windows 11",
        "数据存储：D:\\ssl-research（约 200 GB / 100 GB free）",
    ])

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"saved -> {OUT_PATH}")
    print(f"size  -> {OUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
